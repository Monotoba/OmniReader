from __future__ import annotations

from pathlib import Path
from typing import Any, cast

from ..model import Block, BlockKind, NormalizedDocument
from .base import DocumentParseError, DocumentParser, document_id, make_block


def _paragraph_block(
    paragraph: Any, block_id: str, kind: BlockKind = "paragraph"
) -> Block | None:
    text = paragraph.text.strip()
    if not text:
        return None
    runs = [run for run in paragraph.runs if run.text]
    hidden = bool(runs) and all(run.font.hidden is True for run in runs)
    heuristic = bool(runs) and all(
        (run.font.size is not None and run.font.size.pt < 5)
        or (
            run.font.color.type is not None
            and run.font.color.rgb is not None
            and str(run.font.color.rgb).upper() in {"FFFFFF", "FEFEFE"}
        )
        for run in runs
    )
    if hidden or heuristic:
        kind = "hidden"
    elif paragraph.style and str(paragraph.style.name).lower().startswith("heading"):
        kind = "heading"
    return make_block(block_id, text, kind, metadata={"heuristic": heuristic})


class DocxParser(DocumentParser):
    supported_extensions = (".docx",)

    def parse(self, path: Path) -> NormalizedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentParseError("DOCX support requires python-docx") from exc
        try:
            source = Document(path)
        except Exception as exc:
            raise DocumentParseError(f"Could not read DOCX file: {exc}") from exc
        blocks: list[Block] = []
        for paragraph in source.paragraphs:
            block = _paragraph_block(paragraph, f"body-{len(blocks)}")
            if block:
                blocks.append(block)
        for table in source.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        blocks.append(
                            make_block(f"cell-{len(blocks)}", cell.text, "table_cell")
                        )
        for section in source.sections:
            for part, kind in ((section.header, "header"), (section.footer, "footer")):
                for paragraph in part.paragraphs:
                    block = _paragraph_block(
                        paragraph, f"{kind}-{len(blocks)}", cast(BlockKind, kind)
                    )
                    if block:
                        blocks.append(block)
        title = source.core_properties.title or path.stem
        return NormalizedDocument(
            document_id(path), title, tuple(blocks), {"path": str(path)}
        )
